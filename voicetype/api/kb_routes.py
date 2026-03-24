#!/usr/bin/env python3
# Copyright (C) 2026 VoiceType Contributors
# Licensed under AGPL-3.0

"""
Knowledge base management API.
Supports text entry, file upload (PDF/DOCX/TXT/MD), list, delete, and stats.
"""

import asyncio
import io
import json
import logging
import uuid
from datetime import datetime
from pathlib import Path
from fastapi import APIRouter, HTTPException, UploadFile, File
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

kb_router = APIRouter(prefix="/api/kb")

DATA_DIR = Path("data")
KB_META_FILE = DATA_DIR / "kb_documents.json"

_meta_lock = asyncio.Lock()
_qdrant_lock = asyncio.Lock()


class KBDocument(BaseModel):
    id: str
    title: str
    content: str
    source: str = "text"
    chunks_count: int = 0
    status: str = "pending"
    created_at: str = ""
    error: str = ""


class KBStats(BaseModel):
    total_documents: int = 0
    total_chunks: int = 0
    collection_exists: bool = False
    embedding_model: str = ""


_engine_ref = {"engine": None}


def set_kb_engine(engine):
    _engine_ref["engine"] = engine


async def _load_meta() -> list[dict]:
    async with _meta_lock:
        if KB_META_FILE.exists():
            try:
                with open(KB_META_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                logger.warning("Failed to load KB metadata: %s", e)
                return []
        return []


async def _save_meta(docs: list[dict]):
    async with _meta_lock:
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        with open(KB_META_FILE, "w", encoding="utf-8") as f:
            json.dump(docs, f, ensure_ascii=False, indent=2)


ALLOWED_EXTENSIONS = {".txt", ".md", ".text", ".markdown", ".pdf", ".docx", ".doc"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except ImportError:
        raise HTTPException(500, "PyPDF2 not installed. Run: pip install PyPDF2")


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except ImportError:
        raise HTTPException(500, "python-docx not installed. Run: pip install python-docx")


def _parse_file(file_bytes: bytes, ext: str) -> str:
    """Route file parsing by extension."""
    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_bytes)
    else:
        return file_bytes.decode("utf-8", errors="replace")


@kb_router.get("/documents")
async def list_documents():
    docs = await _load_meta()
    return {"documents": docs, "total": len(docs)}


@kb_router.get("/stats")
async def get_stats():
    docs = await _load_meta()
    completed = [d for d in docs if d.get("status") == "completed"]

    engine = _engine_ref.get("engine")
    embedding_model = ""
    collection_exists = False
    if engine and engine._rag_ext:
        embedding_model = getattr(engine._rag_ext.config, "embedding_model", "BAAI/bge-small-zh-v1.5")
        collection_exists = True

    return KBStats(
        total_documents=len(completed),
        total_chunks=sum(d.get("chunks_count", 0) for d in completed),
        collection_exists=collection_exists,
        embedding_model=embedding_model,
    )


class TextEntry(BaseModel):
    title: str = ""
    content: str


def _task_done_callback(task: asyncio.Task):
    """Log exceptions from fire-and-forget tasks."""
    if task.cancelled():
        return
    exc = task.exception()
    if exc:
        logger.error("Background task %s failed: %s", task.get_name(), exc)


@kb_router.post("/text")
async def add_text(entry: TextEntry):
    """Add a text entry directly to the knowledge base."""
    if not entry.content.strip():
        raise HTTPException(400, "Content cannot be empty")

    doc_id = f"doc_{uuid.uuid4().hex[:12]}"
    title = entry.title or entry.content[:30].strip()

    doc = {
        "id": doc_id,
        "title": title,
        "content": entry.content.strip(),
        "source": "text",
        "chunks_count": 0,
        "status": "pending",
        "created_at": datetime.now().isoformat(),
        "error": "",
    }

    docs = await _load_meta()
    docs.insert(0, doc)
    await _save_meta(docs)

    task = asyncio.create_task(_process_document(doc_id, entry.content.strip()), name=f"kb-process-{doc_id}")
    task.add_done_callback(_task_done_callback)
    return doc


@kb_router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """Upload a document to the knowledge base. Supports PDF, DOCX, TXT, Markdown."""
    if not file.filename:
        raise HTTPException(400, "No file provided")

    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"File type not supported: {ext}. Allowed: PDF, DOCX, TXT, Markdown")

    content_bytes = await file.read()
    if len(content_bytes) > MAX_FILE_SIZE:
        raise HTTPException(400, f"File too large (max {MAX_FILE_SIZE // 1024 // 1024}MB)")

    loop = asyncio.get_event_loop()
    content = await loop.run_in_executor(None, _parse_file, content_bytes, ext)

    if not content.strip():
        raise HTTPException(400, "File content is empty after parsing")

    doc_id = f"doc_{uuid.uuid4().hex[:12]}"

    doc = {
        "id": doc_id,
        "title": file.filename,
        "content": content.strip()[:500],
        "source": "file",
        "file_type": ext,
        "chunks_count": 0,
        "status": "pending",
        "created_at": datetime.now().isoformat(),
        "error": "",
    }

    docs = await _load_meta()
    docs.insert(0, doc)
    await _save_meta(docs)

    task = asyncio.create_task(_process_document(doc_id, content.strip()), name=f"kb-process-{doc_id}")
    task.add_done_callback(_task_done_callback)
    return doc


@kb_router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document from the knowledge base and Qdrant."""
    docs = await _load_meta()
    target = None
    for d in docs:
        if d["id"] == doc_id:
            target = d
            break

    if not target:
        raise HTTPException(404, "Document not found")

    docs = [d for d in docs if d["id"] != doc_id]
    await _save_meta(docs)

    engine = _engine_ref.get("engine")
    if engine and engine._rag_ext:
        try:
            rag = engine._rag_ext
            if hasattr(rag, "_qdrant_client") and rag._qdrant_client:
                from qdrant_client.models import Filter, FieldCondition, MatchValue

                def _do_delete():
                    rag._qdrant_client.delete(
                        collection_name=rag.config.qdrant_collection_name,
                        points_selector=Filter(
                            must=[FieldCondition(key="doc_id", match=MatchValue(value=doc_id))]
                        ),
                    )

                async with _qdrant_lock:
                    await asyncio.get_event_loop().run_in_executor(None, _do_delete)
                logger.info("Deleted vectors for doc %s from Qdrant", doc_id)
        except Exception as e:
            logger.warning("Failed to delete vectors from Qdrant: %s", e)

    return {"status": "ok", "deleted": doc_id}


async def _process_document(doc_id: str, content: str):
    """Split content into chunks and upsert into Qdrant."""
    docs = await _load_meta()
    target = None
    for d in docs:
        if d["id"] == doc_id:
            target = d
            break
    if not target:
        return

    target["status"] = "processing"
    await _save_meta(docs)

    try:
        chunks = _split_text(content)
        if not chunks:
            target["status"] = "completed"
            target["chunks_count"] = 0
            await _save_meta(docs)
            return

        engine = _engine_ref.get("engine")
        if not engine or not engine._rag_ext:
            target["status"] = "completed"
            target["chunks_count"] = len(chunks)
            target["error"] = "RAG not enabled"
            await _save_meta(docs)
            return

        rag = engine._rag_ext
        if not hasattr(rag, "_model") or rag._model is None:
            target["status"] = "failed"
            target["error"] = "Embedding model not loaded"
            await _save_meta(docs)
            return

        from qdrant_client.models import PointStruct, Distance, VectorParams
        import numpy as np

        ENCODE_BATCH = 32
        all_embeddings = []
        for eb_start in range(0, len(chunks), ENCODE_BATCH):
            eb = chunks[eb_start:eb_start + ENCODE_BATCH]
            vecs = rag._model.encode(eb, normalize_embeddings=True)
            all_embeddings.append(vecs)
        embeddings = np.vstack(all_embeddings) if len(all_embeddings) > 1 else all_embeddings[0]
        logger.info("Encoded %d chunks for doc %s, dim=%d", len(chunks), doc_id, embeddings.shape[1])

        collection = rag.config.qdrant_collection_name

        def _qdrant_upsert():
            try:
                rag._qdrant_client.get_collection(collection)
            except Exception:
                rag._qdrant_client.create_collection(
                    collection_name=collection,
                    vectors_config=VectorParams(size=embeddings.shape[1], distance=Distance.COSINE),
                )

            points = []
            for i, (chunk, vec) in enumerate(zip(chunks, embeddings)):
                point_id = str(uuid.uuid4())
                points.append(PointStruct(
                    id=point_id,
                    vector=vec.tolist(),
                    payload={"text": chunk, "doc_id": doc_id, "chunk_index": i},
                ))

            BATCH_SIZE = 50
            for batch_start in range(0, len(points), BATCH_SIZE):
                batch = points[batch_start:batch_start + BATCH_SIZE]
                rag._qdrant_client.upsert(collection_name=collection, points=batch)

            return len(chunks)

        async with _qdrant_lock:
            loop = asyncio.get_event_loop()
            chunk_count = await loop.run_in_executor(None, _qdrant_upsert)

        target["status"] = "completed"
        target["chunks_count"] = chunk_count
        await _save_meta(docs)
        logger.info("Imported %d chunks for doc %s", chunk_count, doc_id)

    except Exception as e:
        logger.error("Failed to process document %s: %s", doc_id, e)
        target["status"] = "failed"
        target["error"] = str(e)
        await _save_meta(docs)


def _split_text(text: str, max_chunk_len: int = 300, overlap: int = 50) -> list[str]:
    """Split text into overlapping chunks by sentence boundaries."""
    import re
    sentences = re.split(r'(?<=[。！？\.\!\?\n])', text)
    sentences = [s.strip() for s in sentences if s.strip()]

    if not sentences:
        return [text] if text.strip() else []

    chunks = []
    current = ""
    for sent in sentences:
        if len(current) + len(sent) > max_chunk_len and current:
            chunks.append(current)
            tail = current[-overlap:] if overlap else ""
            current = tail + sent
        else:
            current += sent

    if current.strip():
        chunks.append(current)

    return chunks
